from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = r"C:\dev\design\参考文献-格式规范版.docx"


references = [
    "[1]World Health Organization.WHO global surveillance and monitoring system for substandard and falsified medical products[R/OL].Geneva:World Health Organization,2017.https://iris.who.int/handle/10665/326708,2017/2026-05-04.",
    "[2]国务院办公厅.国务院办公厅关于加快推进重要产品追溯体系建设的意见:国办发〔2015〕95号[Z/OL].北京:国务院办公厅,2015.https://www.gov.cn/zhengce/zhengceku/2016-01/12/content_10584.htm,2016-01-12/2026-05-04.",
    "[3]国家药品监督管理局.国家药监局关于药品信息化追溯体系建设的指导意见:国药监药管〔2018〕35号[Z/OL].北京:国家药品监督管理局,2018.https://www.gov.cn/zhengce/zhengceku/2018-12/31/content_5434073.htm,2018-11-01/2026-05-04.",
    "[4]全国人民代表大会常务委员会.中华人民共和国药品管理法[Z/OL].北京:全国人民代表大会常务委员会,2019.https://www.gov.cn/xinwen/2019-08/26/content_5424780.htm,2019-08-26/2026-05-04.",
    "[5]NMPAB/T 1001-2019,药品信息化追溯体系建设导则[S].",
    "[6]NMPAB/T 1002-2019,药品追溯码编码要求[S].",
    "[7]NMPAB/T 1003-2019,药品追溯系统基本技术要求[S].",
    "[8]NMPAB/T 1011-2022,药品追溯码标识规范[S].",
    "[9]NMPAB/T 1012-2022,药品追溯消费者查询结果显示规范[S].",
    "[10]唐菀晨,王迎利,张熹,等.我国药品追溯方案和信息化架构研究[J].中国药事,2018,32(7):874-878.",
    "[11]胡泽利.我国药品信息追溯体系建设存在的问题及对策研究[J].中国药房,2019,30(22):3025-3029.",
    "[12]王广平,王颖.我国食品药品安全精准监管实施路径研究[J].中国药事,2019,33(4):355-364.",
    "[13]朱嘉,翁志洁,阮秀芳,等.探讨新修订的《药品管理法》对药品监管工作的挑战和应对策略[J].中国药事,2019,33(12):1335-1340.",
    "[14]U.S.Congress.Drug Quality and Security Act:Public Law 113-54[Z/OL].Washington,DC:U.S.Government Publishing Office,2013.https://www.govinfo.gov/app/details/COMPS-10691,2013/2026-05-04.",
    "[15]European Parliament and the Council.Directive 2011/62/EU[Z/OL].Brussels:Official Journal of the European Union,2011.https://eur-lex.europa.eu/eli/dir/2011/62/oj,2011-07-01/2026-05-04.",
    "[16]European Commission.Commission Delegated Regulation (EU) 2016/161[Z/OL].Brussels:Official Journal of the European Union,2016.https://op.europa.eu/en/publication-detail/-/publication/645fa920-cef8-11e5-a4b5-01aa75ed71a1,2016-02-09/2026-05-04.",
    "[17]Mackey T K,Nayyar G.A review of existing and emerging digital technologies to combat the global trade in fake medicines[J].Expert Opin Drug Saf,2017,16(5):587-602.",
    "[18]张海藩,牟永敏.软件工程导论:第6版[M].北京:清华大学出版社,2013.",
    "[19]王珊,萨师煊.数据库系统概论:第5版[M].北京:高等教育出版社,2014.",
    "[20]黑马程序员.Spring Boot企业级开发教程[M].北京:人民邮电出版社,2019.",
    "[21]李磊.Java EE企业级应用开发实战:Spring Boot+Vue+Element[M].北京:人民邮电出版社,2023.",
    "[22]柳伟卫.Vue.js 3企业级应用开发实战[M].北京:电子工业出版社,2022.",
    "[23]杜文.Flutter实战[M].北京:机械工业出版社,2020.",
    "[24]翟振兴,张恒岩,崔春华,等.深入浅出MySQL:数据库开发、优化与管理维护:第3版[M].北京:人民邮电出版社,2019.",
    "[25]黄健宏.Redis设计与实现[M].北京:机械工业出版社,2014.",
    "[26]朱少民.软件测试方法和技术:第4版[M].北京:清华大学出版社,2022.",
    "[27]Jones M,Bradley J,Sakimura N.JSON Web Token (JWT):RFC 7519[S/OL].Fremont:IETF,2015.https://datatracker.ietf.org/doc/html/rfc7519,2015-05/2026-05-04.",
]


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=10.5, bold=False):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_common(paragraph, first_line=None, hanging=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(18)
    if first_line is not None:
        fmt.first_line_indent = first_line
    if hanging is not None:
        fmt.first_line_indent = -hanging
        fmt.left_indent = hanging


def set_doc_defaults(document):
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    sect_pr = section._sectPr
    doc_grid = sect_pr.find(qn("w:docGrid"))
    if doc_grid is None:
        doc_grid = OxmlElement("w:docGrid")
        sect_pr.append(doc_grid)
    doc_grid.set(qn("w:type"), "lines")
    doc_grid.set(qn("w:linePitch"), "312")


def main():
    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_common(title)
    run = title.add_run("参 考 文 献")
    set_run_font(run, east_asia="黑体", latin="Times New Roman", size=14, bold=True)

    for item in references:
        p = doc.add_paragraph()
        set_paragraph_common(p, hanging=Cm(0.74))
        r = p.add_run(item)
        set_run_font(r, east_asia="宋体", latin="Times New Roman", size=10.5)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
