from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn


OUT = r"C:\dev\design\db_table_structure_appendix_unicode.docx"


def u(text: str) -> str:
    if "\\u" in text:
        return text.encode("ascii").decode("unicode_escape")
    return text


def set_run_font(run, size=12, bold=False):
    run.bold = bold
    run.font.name = u("\\u5b8b\\u4f53")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), u("\\u5b8b\\u4f53"))
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)


def add_para(doc, text, size=12, bold=False, center=False, first_line=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(u(text))
    set_run_font(run, size, bold)
    return p


def set_cell(cell, text, size=10.5, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    run = p.add_run(u(text))
    set_run_font(run, size, bold)


fields = [
    "\\u5b57\\u6bb5\\u540d",
    "\\u6570\\u636e\\u7c7b\\u578b",
    "\\u957f\\u5ea6",
    "\\u952e\\u7c7b\\u578b",
    "\\u662f\\u5426\\u4e3a\\u7a7a",
    "\\u5b57\\u6bb5\\u8bf4\\u660e",
]


tables = [
    (
        "\\u88681 \\u7528\\u6237\\u8868\\uff08user\\uff09\\u7ed3\\u6784\\u8bbe\\u8ba1",
        [
            ["id", "bigint", "20", "\\u4e3b\\u952e", "\\u5426", "\\u7528\\u6237\\u4e3b\\u952eID"],
            ["username", "varchar", "50", "\\u552f\\u4e00", "\\u5426", "\\u7528\\u6237\\u767b\\u5f55\\u540d"],
            ["password", "varchar", "255", "\\u975e\\u952e", "\\u5426", "\\u7528\\u6237\\u767b\\u5f55\\u5bc6\\u7801"],
            ["real_name", "varchar", "50", "\\u975e\\u952e", "\\u5426", "\\u7528\\u6237\\u59d3\\u540d"],
            ["role_id", "bigint", "20", "\\u5916\\u952e", "\\u5426", "\\u89d2\\u8272ID"],
            ["org_name", "varchar", "100", "\\u975e\\u952e", "\\u662f", "\\u6240\\u5c5e\\u7ec4\\u7ec7\\u540d\\u79f0"],
            ["phone", "varchar", "20", "\\u975e\\u952e", "\\u662f", "\\u8054\\u7cfb\\u7535\\u8bdd"],
            ["status", "tinyint", "1", "\\u975e\\u952e", "\\u5426", "\\u8d26\\u53f7\\u72b6\\u6001\\uff0c1\\u542f\\u75280\\u7981\\u7528"],
            ["create_time", "datetime", "-", "\\u975e\\u952e", "\\u5426", "\\u521b\\u5efa\\u65f6\\u95f4"],
            ["update_time", "datetime", "-", "\\u975e\\u952e", "\\u662f", "\\u66f4\\u65b0\\u65f6\\u95f4"],
        ],
    ),
    (
        "\\u88682 \\u89d2\\u8272\\u8868\\uff08role\\uff09\\u7ed3\\u6784\\u8bbe\\u8ba1",
        [
            ["id", "bigint", "20", "\\u4e3b\\u952e", "\\u5426", "\\u89d2\\u8272\\u4e3b\\u952eID"],
            ["role_name", "varchar", "50", "\\u975e\\u952e", "\\u5426", "\\u89d2\\u8272\\u540d\\u79f0"],
            ["role_code", "varchar", "50", "\\u552f\\u4e00", "\\u5426", "\\u89d2\\u8272\\u7f16\\u7801"],
            ["description", "varchar", "255", "\\u975e\\u952e", "\\u662f", "\\u89d2\\u8272\\u63cf\\u8ff0"],
            ["status", "tinyint", "1", "\\u975e\\u952e", "\\u5426", "\\u89d2\\u8272\\u72b6\\u6001"],
            ["create_time", "datetime", "-", "\\u975e\\u952e", "\\u5426", "\\u521b\\u5efa\\u65f6\\u95f4"],
            ["update_time", "datetime", "-", "\\u975e\\u952e", "\\u662f", "\\u66f4\\u65b0\\u65f6\\u95f4"],
        ],
    ),
    (
        "\\u88683 \\u836f\\u54c1\\u4fe1\\u606f\\u8868\\uff08drug_info\\uff09\\u7ed3\\u6784\\u8bbe\\u8ba1",
        [
            ["id", "bigint", "20", "\\u4e3b\\u952e", "\\u5426", "\\u836f\\u54c1\\u4e3b\\u952eID"],
            ["drug_code", "varchar", "50", "\\u552f\\u4e00", "\\u5426", "\\u836f\\u54c1\\u7f16\\u7801"],
            ["drug_name", "varchar", "100", "\\u975e\\u952e", "\\u5426", "\\u836f\\u54c1\\u540d\\u79f0"],
            ["specification", "varchar", "100", "\\u975e\\u952e", "\\u662f", "\\u89c4\\u683c\\u578b\\u53f7"],
            ["manufacturer", "varchar", "100", "\\u975e\\u952e", "\\u662f", "\\u751f\\u4ea7\\u4f01\\u4e1a"],
            ["approval_no", "varchar", "100", "\\u975e\\u952e", "\\u662f", "\\u6279\\u51c6\\u6587\\u53f7"],
            ["category", "varchar", "50", "\\u975e\\u952e", "\\u662f", "\\u836f\\u54c1\\u5206\\u7c7b"],
            ["status", "tinyint", "1", "\\u975e\\u952e", "\\u5426", "\\u836f\\u54c1\\u72b6\\u6001"],
            ["remark", "varchar", "255", "\\u975e\\u952e", "\\u662f", "\\u5907\\u6ce8"],
            ["create_time", "datetime", "-", "\\u975e\\u952e", "\\u5426", "\\u521b\\u5efa\\u65f6\\u95f4"],
        ],
    ),
    (
        "\\u88684 \\u751f\\u4ea7\\u6279\\u6b21\\u8868\\uff08production_batch\\uff09\\u7ed3\\u6784\\u8bbe\\u8ba1",
        [
            ["id", "bigint", "20", "\\u4e3b\\u952e", "\\u5426", "\\u6279\\u6b21\\u4e3b\\u952eID"],
            ["batch_no", "varchar", "50", "\\u552f\\u4e00", "\\u5426", "\\u751f\\u4ea7\\u6279\\u6b21\\u53f7"],
            ["drug_id", "bigint", "20", "\\u5916\\u952e", "\\u5426", "\\u5173\\u8054\\u836f\\u54c1ID"],
            ["production_date", "date", "-", "\\u975e\\u952e", "\\u5426", "\\u751f\\u4ea7\\u65e5\\u671f"],
            ["expiry_date", "date", "-", "\\u975e\\u952e", "\\u5426", "\\u6709\\u6548\\u671f\\u81f3"],
            ["quantity", "int", "11", "\\u975e\\u952e", "\\u5426", "\\u751f\\u4ea7\\u6570\\u91cf"],
            ["producer", "varchar", "100", "\\u975e\\u952e", "\\u662f", "\\u751f\\u4ea7\\u4e3b\\u4f53"],
            ["status", "tinyint", "1", "\\u975e\\u952e", "\\u5426", "\\u6279\\u6b21\\u72b6\\u6001"],
            ["create_time", "datetime", "-", "\\u975e\\u952e", "\\u5426", "\\u521b\\u5efa\\u65f6\\u95f4"],
        ],
    ),
    (
        "\\u88685 \\u91c7\\u8d2d\\u8bb0\\u5f55\\u8868\\uff08procurement_record\\uff09\\u7ed3\\u6784\\u8bbe\\u8ba1",
        [
            ["id", "bigint", "20", "\\u4e3b\\u952e", "\\u5426", "\\u91c7\\u8d2d\\u8bb0\\u5f55ID"],
            ["batch_id", "bigint", "20", "\\u5916\\u952e", "\\u5426", "\\u5173\\u8054\\u6279\\u6b21ID"],
            ["drug_id", "bigint", "20", "\\u5916\\u952e", "\\u5426", "\\u836f\\u54c1ID"],
            ["supplier_name", "varchar", "100", "\\u975e\\u952e", "\\u662f", "\\u4f9b\\u5e94\\u5546\\u540d\\u79f0"],
            ["purchaser_name", "varchar", "100", "\\u975e\\u952e", "\\u662f", "\\u91c7\\u8d2d\\u65b9\\u540d\\u79f0"],
            ["quantity", "int", "11", "\\u975e\\u952e", "\\u5426", "\\u91c7\\u8d2d\\u6570\\u91cf"],
            ["procurement_time", "datetime", "-", "\\u975e\\u952e", "\\u5426", "\\u91c7\\u8d2d\\u65f6\\u95f4"],
            ["remark", "varchar", "255", "\\u975e\\u952e", "\\u662f", "\\u5907\\u6ce8"],
        ],
    ),
    (
        "\\u88686 \\u9500\\u552e\\u8bb0\\u5f55\\u8868\\uff08sale_record\\uff09\\u7ed3\\u6784\\u8bbe\\u8ba1",
        [
            ["id", "bigint", "20", "\\u4e3b\\u952e", "\\u5426", "\\u9500\\u552e\\u8bb0\\u5f55ID"],
            ["batch_id", "bigint", "20", "\\u5916\\u952e", "\\u5426", "\\u5173\\u8054\\u6279\\u6b21ID"],
            ["drug_id", "bigint", "20", "\\u5916\\u952e", "\\u5426", "\\u836f\\u54c1ID"],
            ["customer_name", "varchar", "100", "\\u975e\\u952e", "\\u662f", "\\u8d2d\\u4e70\\u65b9\\u540d\\u79f0"],
            ["quantity", "int", "11", "\\u975e\\u952e", "\\u5426", "\\u9500\\u552e\\u6570\\u91cf"],
            ["sale_time", "datetime", "-", "\\u975e\\u952e", "\\u5426", "\\u9500\\u552e\\u65f6\\u95f4"],
            ["operator_name", "varchar", "50", "\\u975e\\u952e", "\\u662f", "\\u64cd\\u4f5c\\u4eba"],
            ["remark", "varchar", "255", "\\u975e\\u952e", "\\u662f", "\\u5907\\u6ce8"],
        ],
    ),
    (
        "\\u88687 \\u5e93\\u5b58\\u8868\\uff08inventory\\uff09\\u7ed3\\u6784\\u8bbe\\u8ba1",
        [
            ["id", "bigint", "20", "\\u4e3b\\u952e", "\\u5426", "\\u5e93\\u5b58\\u8bb0\\u5f55ID"],
            ["drug_id", "bigint", "20", "\\u5916\\u952e", "\\u5426", "\\u836f\\u54c1ID"],
            ["batch_id", "bigint", "20", "\\u5916\\u952e", "\\u5426", "\\u6279\\u6b21ID"],
            ["org_name", "varchar", "100", "\\u975e\\u952e", "\\u5426", "\\u6240\\u5c5e\\u7ec4\\u7ec7"],
            ["stock_quantity", "int", "11", "\\u975e\\u952e", "\\u5426", "\\u5f53\\u524d\\u5e93\\u5b58\\u6570\\u91cf"],
            ["update_time", "datetime", "-", "\\u975e\\u952e", "\\u5426", "\\u5e93\\u5b58\\u66f4\\u65b0\\u65f6\\u95f4"],
            ["status", "tinyint", "1", "\\u975e\\u952e", "\\u5426", "\\u5e93\\u5b58\\u72b6\\u6001"],
        ],
    ),
    (
        "\\u88688 \\u76d1\\u7ba1\\u4efb\\u52a1\\u8868\\uff08regulatory_task\\uff09\\u7ed3\\u6784\\u8bbe\\u8ba1",
        [
            ["id", "bigint", "20", "\\u4e3b\\u952e", "\\u5426", "\\u76d1\\u7ba1\\u4efb\\u52a1ID"],
            ["task_title", "varchar", "100", "\\u975e\\u952e", "\\u5426", "\\u4efb\\u52a1\\u6807\\u9898"],
            ["clue_source", "varchar", "100", "\\u975e\\u952e", "\\u662f", "\\u7ebf\\u7d22\\u6765\\u6e90"],
            ["problem_desc", "varchar", "255", "\\u975e\\u952e", "\\u5426", "\\u95ee\\u9898\\u63cf\\u8ff0"],
            ["severity_level", "varchar", "20", "\\u975e\\u952e", "\\u662f", "\\u4e25\\u91cd\\u7a0b\\u5ea6"],
            ["assigned_to", "varchar", "50", "\\u975e\\u952e", "\\u662f", "\\u6d3e\\u53d1\\u5bf9\\u8c61"],
            ["status", "varchar", "20", "\\u975e\\u952e", "\\u5426", "\\u4efb\\u52a1\\u72b6\\u6001"],
            ["result_desc", "varchar", "255", "\\u975e\\u952e", "\\u662f", "\\u5904\\u7406\\u7ed3\\u679c\\u8bf4\\u660e"],
            ["create_time", "datetime", "-", "\\u975e\\u952e", "\\u5426", "\\u521b\\u5efa\\u65f6\\u95f4"],
            ["finish_time", "datetime", "-", "\\u975e\\u952e", "\\u662f", "\\u5b8c\\u6210\\u65f6\\u95f4"],
        ],
    ),
    (
        "\\u88689 \\u4e3e\\u62a5/\\u4e0d\\u826f\\u53cd\\u5e94\\u8868\\uff08adverse_reaction\\uff09\\u7ed3\\u6784\\u8bbe\\u8ba1",
        [
            ["id", "bigint", "20", "\\u4e3b\\u952e", "\\u5426", "\\u4e3e\\u62a5\\u8bb0\\u5f55ID"],
            ["reporter_name", "varchar", "50", "\\u975e\\u952e", "\\u662f", "\\u4e3e\\u62a5\\u4eba\\u59d3\\u540d"],
            ["phone", "varchar", "20", "\\u975e\\u952e", "\\u662f", "\\u8054\\u7cfb\\u7535\\u8bdd"],
            ["drug_name", "varchar", "100", "\\u975e\\u952e", "\\u5426", "\\u6d89\\u53ca\\u836f\\u54c1\\u540d\\u79f0"],
            ["description", "varchar", "255", "\\u975e\\u952e", "\\u5426", "\\u95ee\\u9898\\u63cf\\u8ff0"],
            ["report_time", "datetime", "-", "\\u975e\\u952e", "\\u5426", "\\u4e3e\\u62a5\\u65f6\\u95f4"],
            ["status", "varchar", "20", "\\u975e\\u952e", "\\u5426", "\\u5904\\u7406\\u72b6\\u6001"],
        ],
    ),
    (
        "\\u886810 \\u6d88\\u606f\\u901a\\u77e5\\u8868\\uff08message_notice\\uff09\\u7ed3\\u6784\\u8bbe\\u8ba1",
        [
            ["id", "bigint", "20", "\\u4e3b\\u952e", "\\u5426", "\\u6d88\\u606fID"],
            ["title", "varchar", "100", "\\u975e\\u952e", "\\u5426", "\\u6d88\\u606f\\u6807\\u9898"],
            ["content", "varchar", "255", "\\u975e\\u952e", "\\u5426", "\\u6d88\\u606f\\u5185\\u5bb9"],
            ["receiver_id", "bigint", "20", "\\u5916\\u952e", "\\u5426", "\\u63a5\\u6536\\u7528\\u6237ID"],
            ["read_status", "tinyint", "1", "\\u975e\\u952e", "\\u5426", "\\u5df2\\u8bfb\\u72b6\\u6001"],
            ["send_time", "datetime", "-", "\\u975e\\u952e", "\\u5426", "\\u53d1\\u9001\\u65f6\\u95f4"],
        ],
    ),
]


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

add_para(doc, "\\u6570\\u636e\\u5e93\\u8868\\u7ed3\\u6784\\u8bbe\\u8ba1\\u8865\\u5145", size=16, bold=True, center=True)
add_para(doc, "\\u4e3a\\u4fbf\\u4e8e\\u8bf4\\u660e\\u7cfb\\u7edf\\u6570\\u636e\\u5e93\\u8bbe\\u8ba1\\u5185\\u5bb9\\uff0c\\u672c\\u6587\\u5bf9\\u4e3b\\u8981\\u4e1a\\u52a1\\u6570\\u636e\\u8868\\u7684\\u7ed3\\u6784\\u8fdb\\u884c\\u8868\\u683c\\u5316\\u63cf\\u8ff0\\u3002\\u5404\\u8868\\u56f4\\u7ed5\\u6743\\u9650\\u7ba1\\u7406\\u3001\\u836f\\u54c1\\u57fa\\u7840\\u4fe1\\u606f\\u3001\\u6d41\\u901a\\u8ffd\\u6eaf\\u548c\\u76d1\\u7ba1\\u4e1a\\u52a1\\u7b49\\u529f\\u80fd\\u6a21\\u5757\\u8fdb\\u884c\\u8bbe\\u8ba1\\uff0c\\u5b57\\u6bb5\\u8bbe\\u7f6e\\u517c\\u987e\\u4e1a\\u52a1\\u5b8c\\u6574\\u6027\\u3001\\u6570\\u636e\\u4e00\\u81f4\\u6027\\u548c\\u540e\\u7eed\\u67e5\\u8be2\\u7edf\\u8ba1\\u9700\\u6c42\\u3002\\u4e0b\\u9762\\u9009\\u53d6\\u7cfb\\u7edf\\u4e2d\\u5177\\u6709\\u4ee3\\u8868\\u6027\\u7684\\u6838\\u5fc3\\u6570\\u636e\\u8868\\u8fdb\\u884c\\u8bf4\\u660e\\u3002", first_line=True)
add_para(doc, "\\u8868\\u683c\\u5b57\\u6bb5\\u7edf\\u4e00\\u5305\\u62ec\\u5b57\\u6bb5\\u540d\\u3001\\u6570\\u636e\\u7c7b\\u578b\\u3001\\u957f\\u5ea6\\u3001\\u952e\\u7c7b\\u578b\\u3001\\u662f\\u5426\\u4e3a\\u7a7a\\u548c\\u5b57\\u6bb5\\u8bf4\\u660e\\uff0c\\u4ee5\\u4fbf\\u5728\\u8bba\\u6587\\u4e2d\\u76f4\\u63a5\\u4f7f\\u7528\\u6216\\u8fdb\\u4e00\\u6b65\\u8c03\\u6574\\u3002", first_line=True)

for title, rows in tables:
    add_para(doc, title, size=12, bold=True, center=True)
    table = doc.add_table(rows=1, cols=len(fields))
    table.style = "Table Grid"
    for i, f in enumerate(fields):
        set_cell(table.rows[0].cells[i], f, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    doc.add_paragraph("")

add_para(doc, "\\u901a\\u8fc7\\u4e0a\\u8ff0\\u6570\\u636e\\u8868\\u8bbe\\u8ba1\\u53ef\\u4ee5\\u770b\\u51fa\\uff0c\\u672c\\u7cfb\\u7edf\\u6570\\u636e\\u5e93\\u7ed3\\u6784\\u80fd\\u591f\\u8f83\\u597d\\u652f\\u6301\\u7528\\u6237\\u6743\\u9650\\u7ba1\\u7406\\u3001\\u836f\\u54c1\\u57fa\\u7840\\u4fe1\\u606f\\u7ef4\\u62a4\\u3001\\u6d41\\u901a\\u8bb0\\u5f55\\u8ffd\\u8e2a\\u3001\\u76d1\\u7ba1\\u4efb\\u52a1\\u5904\\u7f6e\\u4ee5\\u53ca\\u6d88\\u606f\\u901a\\u77e5\\u7b49\\u4e1a\\u52a1\\u9700\\u6c42\\u3002\\u5404\\u6570\\u636e\\u8868\\u4e4b\\u95f4\\u901a\\u8fc7\\u4e3b\\u952e\\u548c\\u5916\\u952e\\u5efa\\u7acb\\u5173\\u8054\\u5173\\u7cfb\\uff0c\\u65e2\\u4fdd\\u8bc1\\u4e86\\u6570\\u636e\\u5b58\\u50a8\\u7684\\u89c4\\u8303\\u6027\\uff0c\\u4e5f\\u4e3a\\u540e\\u7eed\\u8ffd\\u6eaf\\u67e5\\u8be2\\u3001\\u7edf\\u8ba1\\u5206\\u6790\\u548c\\u76d1\\u7ba1\\u95ed\\u73af\\u5904\\u7406\\u63d0\\u4f9b\\u4e86\\u53ef\\u9760\\u7684\\u6570\\u636e\\u652f\\u6491\\u3002", first_line=True)

doc.save(OUT)
print(OUT)
